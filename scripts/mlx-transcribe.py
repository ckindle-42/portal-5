#!/usr/bin/env python3
"""
MLX Transcribe Server — Portal 5 (TASK-TRANSCRIBE-001)

Host-native transcription server for Apple Silicon.
- Fast ASR: Parakeet-TDT-v3 (Metal-accelerated, word-level timestamps)
- Diarized: Parakeet transcript + Sortformer speaker diarization, merged at the
  word level. Parakeet always produces the full transcript; Sortformer only adds
  "who spoke when". A monologue → one speaker; a conversation → labels per turn.
  If diarization fails or the file is very long, the transcript still comes back
  (single-speaker) rather than truncating. No HuggingFace token required.
- Output: every transcription (plain or diarized) writes three sidecars into the
  workspace generated/transcripts/ directory — JSON (canonical), Markdown, and
  Word (.docx). The .docx is skipped only if python-docx is unavailable.

Runs on the host (not Docker) — same pattern as mlx-proxy.py and mlx-speech.py.
Open WebUI / Pipeline connects via host.docker.internal:8924.

Usage:
    python scripts/mlx-transcribe.py
    # or via launch.sh:
    ./launch.sh start-transcribe

Files reachable:
- OWUI uploads: ${AI_OUTPUT_DIR}/uploads/<file_id> (resolved via workspace helper)
- Outputs:      ${AI_OUTPUT_DIR}/generated/transcripts/transcript_<uuid>.{json,md,docx}
"""

from __future__ import annotations

import asyncio
import contextlib
import json
import logging
import os
import re
import sys
import tempfile
import time
import uuid
from pathlib import Path
from typing import Any

import uvicorn
from fastapi import FastAPI, File, Form, Request, UploadFile
from fastapi.responses import FileResponse, JSONResponse

# Workspace helpers — TASK-WORKSPACE-001
# Add the repo root to sys.path so this host-native script can import portal.platform.mcp_host
_REPO_ROOT = Path(__file__).resolve().parent.parent
if str(_REPO_ROOT) not in sys.path:
    sys.path.insert(0, str(_REPO_ROOT))

from portal.platform.mcp_host.workspace import (  # noqa: E402
    get_generated_dir,
    get_uploads_dir,
    resolve_upload_path,
)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(name)s] %(levelname)s: %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger("mlx-transcribe")

# ── Configuration ──────────────────────────────────────────────────────────────

PORT = int(os.getenv("MLX_TRANSCRIBE_PORT", "8924"))
HOST = os.getenv("MLX_TRANSCRIBE_HOST", "0.0.0.0")

# ASR (transcript + word-level timestamps) — Parakeet-TDT-v3. Used by both tools.
PARAKEET_MODEL = os.getenv("MLX_PARAKEET_MODEL", "mlx-community/parakeet-tdt-0.6b-v3")
# Parakeet's conformer encoder is O(n^2) in sequence length — a single forward
# pass OOMs somewhere past ~20 min of audio. Chunk long files (mlx-audio merges
# the per-chunk tokens across the overlap). ~40x realtime, so chunking is cheap.
PARAKEET_CHUNK_S = float(os.getenv("MLX_PARAKEET_CHUNK_S", "120"))
PARAKEET_OVERLAP_S = float(os.getenv("MLX_PARAKEET_OVERLAP_S", "15"))
# Speaker diarization ("who spoke when") — NVIDIA Sortformer, MLX port. 4-speaker
# ceiling; runs in one forward pass (no HF token, no pyannote).
DIARIZE_MODEL = os.getenv("MLX_DIARIZE_MODEL", "mlx-community/diar_sortformer_4spk-v1-fp32")
# Safety valve: above this duration (s) diarization is skipped and the transcript
# is returned single-speaker with a warning. Streaming diarization has bounded
# memory, so this is really an OWUI-wall-clock guard, not a memory one.
DIARIZE_MAX_S = float(os.getenv("MLX_DIARIZE_MAX_S", "10800"))

# ── Model cache ────────────────────────────────────────────────────────────────

_parakeet_model: Any = None
_diarizer_model: Any = None
_pipeline_lock = asyncio.Semaphore(1)  # GPU-heavy; serialize.


def _get_parakeet() -> Any:
    """Lazy-load and cache Parakeet-TDT-v3 (transcript + word timestamps)."""
    global _parakeet_model
    if _parakeet_model is None:
        from mlx_audio.stt.utils import load

        logger.info("Loading Parakeet ASR: %s", PARAKEET_MODEL)
        _parakeet_model = load(PARAKEET_MODEL)
        logger.info("Parakeet ready")
    return _parakeet_model


def _get_diarizer() -> Any:
    """Lazy-load and cache the Sortformer speaker-diarization model."""
    global _diarizer_model
    if _diarizer_model is None:
        from mlx_audio.vad import load

        logger.info("Loading Sortformer diarizer: %s", DIARIZE_MODEL)
        _diarizer_model = load(DIARIZE_MODEL)
        logger.info("Sortformer ready")
    return _diarizer_model


# ── Core pipeline ──────────────────────────────────────────────────────────────


def _audio_seconds(audio_path: str) -> float:
    """Best-effort audio duration in seconds (0.0 if it can't be read cheaply)."""
    try:
        import soundfile as sf

        return float(sf.info(audio_path).duration)
    except Exception:  # noqa: BLE001
        pass
    try:
        import wave

        with wave.open(audio_path, "rb") as w:
            return w.getnframes() / float(w.getframerate())
    except Exception:  # noqa: BLE001
        return 0.0


def _parakeet_generate(audio_path: str) -> Any:
    """Parakeet transcription, chunked when the file is long enough to OOM a
    single forward pass. mlx-audio merges the per-chunk tokens across the overlap.
    A single pass that OOMs anyway is retried chunked.
    """
    model = _get_parakeet()
    if _audio_seconds(audio_path) > PARAKEET_CHUNK_S:
        return model.generate(
            audio_path,
            chunk_duration=PARAKEET_CHUNK_S,
            overlap_duration=PARAKEET_OVERLAP_S,
        )
    try:
        return model.generate(audio_path)
    except RuntimeError as e:
        if "malloc" not in str(e) and "buffer size" not in str(e):
            raise
        logger.warning("Parakeet single pass OOMed (%s) — retrying chunked", e)
        return model.generate(
            audio_path,
            chunk_duration=PARAKEET_CHUNK_S,
            overlap_duration=PARAKEET_OVERLAP_S,
        )


def _parakeet_words(audio_path: str) -> tuple[str, list[dict], float]:
    """Run Parakeet and return (full_text, words, duration).

    ``words`` is a flat list of {start, end, text} at token granularity — Parakeet
    tokens are SentencePiece word-pieces whose ``text`` carries a leading space at
    word starts, so ``"".join(w["text"] ...)`` reconstructs spacing exactly.
    """
    result = _parakeet_generate(audio_path)
    words: list[dict] = []
    for sent in getattr(result, "sentences", []):
        for tok in getattr(sent, "tokens", []):
            words.append({"start": float(tok.start), "end": float(tok.end), "text": tok.text})
    duration = words[-1]["end"] if words else 0.0
    return result.text.strip(), words, round(duration, 2)


def _transcribe(audio_path: str, language: str | None) -> dict:
    """Fast ASR via Parakeet-TDT-v3 (no speaker labels).

    Returns {text, language, duration, segments} where segments are sentences.
    """
    result = _parakeet_generate(audio_path)
    segments = [
        {"start": round(float(s.start), 2), "end": round(float(s.end), 2), "text": s.text.strip()}
        for s in getattr(result, "sentences", [])
    ]
    duration = segments[-1]["end"] if segments else 0.0
    return {
        "text": result.text.strip(),
        "language": language or "en",
        "duration": round(duration, 2),
        "segments": segments,
    }


# Sortformer's full-context forward pass gives the best speaker consistency but
# its conformer attention is O(n^2): measured clean to 20 min on 64 GB, thrashing
# swap by 25 min, hard OOM by 40 min. Below this cutoff use it; above it fall back
# to the streaming path (fixed chunks + a speaker cache — bounded memory, slightly
# noisier speaker identity, still fine for a long recording).
DIARIZE_FULL_CONTEXT_MAX_S = float(os.getenv("MLX_DIARIZE_FULL_CONTEXT_MAX_S", "900"))


def _diarize(audio_path: str) -> list[dict]:
    """Sortformer speaker turns → [{start, end, speaker:int}] (raw speaker ids 0-3)."""
    model = _get_diarizer()
    threshold = float(os.getenv("MLX_DIARIZE_THRESHOLD", "0.5"))
    min_duration = float(os.getenv("MLX_DIARIZE_MIN_DURATION", "0.3"))
    merge_gap = float(os.getenv("MLX_DIARIZE_MERGE_GAP", "0.5"))

    raw: list[dict] = []
    if _audio_seconds(audio_path) <= DIARIZE_FULL_CONTEXT_MAX_S:
        out = model.generate(
            audio_path, threshold=threshold, min_duration=min_duration, merge_gap=merge_gap
        )
        raw = [
            {"start": float(s.start), "end": float(s.end), "speaker": int(s.speaker)}
            for s in out.segments
        ]
    else:
        chunk_s = float(os.getenv("MLX_DIARIZE_CHUNK_S", "20"))
        for chunk in model.generate_stream(
            audio_path,
            chunk_duration=chunk_s,
            threshold=threshold,
            min_duration=min_duration,
            merge_gap=merge_gap,
        ):
            raw.extend(
                {"start": float(s.start), "end": float(s.end), "speaker": int(s.speaker)}
                for s in chunk.segments
            )

    # Stitch chunk-boundary splits: adjacent same-speaker segments within merge_gap.
    raw.sort(key=lambda s: s["start"])
    merged: list[dict] = []
    for s in raw:
        if (
            merged
            and merged[-1]["speaker"] == s["speaker"]
            and s["start"] - merged[-1]["end"] <= merge_gap
        ):
            merged[-1]["end"] = max(merged[-1]["end"], s["end"])
        else:
            merged.append(dict(s))
    return merged


def _speaker_for_span(diar: list[dict], w0: float, w1: float) -> int | None:
    """Speaker whose turns overlap [w0, w1] most; fall back to the nearest turn."""
    if not diar:
        return None
    per: dict[int, float] = {}
    for seg in diar:
        ov = max(0.0, min(w1, seg["end"]) - max(w0, seg["start"]))
        if ov > 0:
            per[seg["speaker"]] = per.get(seg["speaker"], 0.0) + ov
    if per:
        return max(per, key=lambda k: per[k])
    mid = (w0 + w1) / 2
    nearest = min(diar, key=lambda s: min(abs(s["start"] - mid), abs(s["end"] - mid)))
    return nearest["speaker"]


def _smooth_speaker_runs(raw: list[int], words: list[dict]) -> list[int]:
    """Flip a short speaker run wedged between two runs of the same other speaker.

    Such a run is almost always a diarization boundary wobble (e.g. "we'"/"re"
    split across a turn edge), not a real one-word interjection.
    """
    min_turn_s = float(os.getenv("MLX_DIARIZE_MIN_TURN", "1.0"))
    out = list(raw)
    i = 0
    while i < len(out):
        j = i
        while j < len(out) and out[j] == out[i]:
            j += 1
        span = words[j - 1]["end"] - words[i]["start"]
        prev_spk = out[i - 1] if i > 0 else None
        next_spk = out[j] if j < len(out) else None
        if span < min_turn_s and prev_spk is not None and prev_spk == next_spk:
            for k in range(i, j):
                out[k] = prev_spk
        i = j
    return out


def _merge_words_and_speakers(
    words: list[dict], diar: list[dict], num_speakers: int | None
) -> tuple[list[dict], int]:
    """Assign every Parakeet word to a Sortformer speaker, then group into turns.

    - Each word goes to the speaker with the most temporal overlap (nearest turn
      if none overlaps) — so a speaker change lands on a word boundary, never
      mid-word.
    - If ``num_speakers`` is given and diarization found more, the least-talkative
      raw speakers are folded into the nearest kept speaker.
    - Kept speakers are renumbered by order of first appearance, so ``SPEAKER_00``
      is whoever speaks first.
    """
    if not words:
        return [], 0

    raw = [_speaker_for_span(diar, w["start"], w["end"]) for w in words]
    raw = [r if r is not None else 0 for r in raw]

    # Sentence-final punctuation is emitted as its own token timestamped right on
    # the turn boundary — pin it to the previous word's speaker so it never opens
    # or closes a turn on its own (" ? Good." → "? " stays with the question).
    for i in range(1, len(words)):
        if words[i]["text"].strip() in {".", ",", "?", "!", ";", ":", "…", "-", "—"}:
            raw[i] = raw[i - 1]

    # Optional cap: keep the N speakers with the most words, remap the rest.
    if num_speakers and num_speakers > 0:
        totals: dict[int, float] = {}
        for spk, w in zip(raw, words, strict=True):
            totals[spk] = totals.get(spk, 0.0) + (w["end"] - w["start"])
        if len(totals) > num_speakers:
            keep = {s for s, _ in sorted(totals.items(), key=lambda kv: -kv[1])[:num_speakers]}
            kept_turns = [s for s in diar if s["speaker"] in keep] or diar
            raw = [
                r if r in keep else (_speaker_for_span(kept_turns, w["start"], w["end"]) or 0)
                for r, w in zip(raw, words, strict=True)
            ]

    raw = _smooth_speaker_runs(raw, words)

    # Renumber by first appearance.
    order: dict[int, int] = {}
    for r in raw:
        if r not in order:
            order[r] = len(order)
    final = [order[r] for r in raw]

    # Group consecutive same-speaker words into turns.
    turns: list[dict] = []
    for spk, w in zip(final, words, strict=True):
        label = f"SPEAKER_{spk:02d}"
        if turns and turns[-1]["speaker"] == label:
            turns[-1]["end"] = round(w["end"], 2)
            turns[-1]["_text"] += w["text"]
        else:
            turns.append(
                {
                    "start": round(w["start"], 2),
                    "end": round(w["end"], 2),
                    "speaker": label,
                    "_text": w["text"],
                }
            )
    segments = [
        {"start": t["start"], "end": t["end"], "speaker": t["speaker"], "text": t["_text"].strip()}
        for t in turns
    ]
    return segments, len(order)


def _diarized_transcribe(audio_path: str, num_speakers: int | None, language: str | None) -> dict:
    """Parakeet transcript + Sortformer diarization, merged at word level.

    The transcript is always complete. If diarization is skipped (file too long) or
    fails, everything is returned as a single speaker with a ``warning``.
    Returns {text, language, duration, speaker_count, segments, diarize_s, warning?}.
    """
    text, words, duration = _parakeet_words(audio_path)

    diar: list[dict] = []
    diarize_s = 0.0
    warning: str | None = None
    if duration > DIARIZE_MAX_S:
        warning = (
            f"audio is {duration / 60:.0f} min — past the {DIARIZE_MAX_S / 60:.0f} min "
            f"single-pass diarization ceiling; transcript returned without speaker labels"
        )
        logger.warning("diarization skipped: %s", warning)
    else:
        t0 = time.time()
        try:
            diar = _diarize(audio_path)
            diarize_s = round(time.time() - t0, 2)
        except Exception as e:  # noqa: BLE001 — any diarizer failure degrades gracefully
            warning = (
                f"speaker diarization failed ({e}); transcript returned without speaker labels"
            )
            logger.warning("diarization failed: %s", e, exc_info=True)

    segments, speaker_count = _merge_words_and_speakers(words, diar, num_speakers)
    if not diar and segments:
        speaker_count = 1

    out: dict = {
        "text": text,
        "language": language or "en",
        "duration": duration,
        "speaker_count": speaker_count,
        "segments": segments,
        "diarize_s": diarize_s,
    }
    if warning:
        out["warning"] = warning
    return out


def _format_markdown(segments: list[dict], meta: dict, source_name: str = "audio") -> str:
    """Render transcript segments as markdown. Speaker labels are emitted when a
    segment carries a ``speaker`` key (diarized path); otherwise it is a plain
    timestamped transcript."""
    lines = [
        f"# Transcript: {source_name}",
        "",
        f"- **Duration**: {meta['duration']:.1f}s",
        f"- **Language**: {meta['language']}",
    ]
    if meta.get("speaker_count") is not None:
        lines.append(f"- **Speakers**: {meta['speaker_count']}")
    lines += [
        f"- **Generated**: {time.strftime('%Y-%m-%d %H:%M:%S')}",
        "",
        "---",
        "",
    ]
    for seg in segments:
        ts = f"[{int(seg['start']) // 60:02d}:{int(seg['start']) % 60:02d}]"
        lines.append(f"**{seg['speaker']}** {ts}" if seg.get("speaker") else ts)
        lines.append(seg["text"])
        lines.append("")
    return "\n".join(lines)


def _write_docx(segments: list[dict], meta: dict, source_name: str, path: Path) -> bool:
    """Write the transcript as a Word document. Returns False and writes nothing
    if python-docx is unavailable — the JSON and Markdown sidecars still stand."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        logger.warning("python-docx not installed — skipping .docx sidecar")
        return False

    doc = Document()
    doc.add_heading(f"Transcript: {source_name}", level=1)
    meta_bits = [f"Duration: {meta['duration']:.1f}s", f"Language: {meta['language']}"]
    if meta.get("speaker_count") is not None:
        meta_bits.append(f"Speakers: {meta['speaker_count']}")
    meta_bits.append(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph().add_run("  ·  ".join(meta_bits)).italic = True

    for seg in segments:
        ts = f"[{int(seg['start']) // 60:02d}:{int(seg['start']) % 60:02d}]"
        para = doc.add_paragraph()
        label = para.add_run((f"{seg['speaker']} {ts}" if seg.get("speaker") else ts) + "  ")
        label.bold = True
        label.font.size = Pt(9)
        para.add_run(seg["text"])

    doc.save(str(path))
    return True


def _persist_transcript(segments: list[dict], meta: dict, source_name: str) -> dict:
    """Write JSON + Markdown + Word sidecars into the shared workspace and return
    the paths, download URLs and rendered markdown. Both the plain and the
    diarized paths route through here so every transcript yields the same three
    artifacts."""
    out_dir = get_generated_dir("transcripts")
    uid = uuid.uuid4().hex[:12]
    json_path = out_dir / f"transcript_{uid}.json"
    md_path = out_dir / f"transcript_{uid}.md"
    docx_path = out_dir / f"transcript_{uid}.docx"

    json_path.write_text(
        json.dumps({**meta, "segments": segments, "source": source_name}, indent=2)
    )
    markdown = _format_markdown(segments, meta, source_name)
    md_path.write_text(markdown)
    docx_written = _write_docx(segments, meta, source_name, docx_path)

    result = {
        **meta,
        "segments": segments,
        "markdown": markdown,
        "json_path": str(json_path),
        "md_path": str(md_path),
        "json_url": f"http://host.docker.internal:{PORT}/files/{json_path.name}",
        "md_url": f"http://host.docker.internal:{PORT}/files/{md_path.name}",
    }
    if docx_written:
        result["docx_path"] = str(docx_path)
        result["docx_url"] = f"http://host.docker.internal:{PORT}/files/{docx_path.name}"
    else:
        result["warning"] = (
            f"{meta.get('warning', '')} (.docx sidecar skipped: python-docx unavailable)".strip()
        )
    return result


def _run_pipeline(
    audio_path: str,
    language: str | None,
    num_speakers: int | None,
    source_name: str = "audio",
) -> dict:
    """Synchronous diarized pipeline (Parakeet + Sortformer). Caller wraps in asyncio.to_thread."""
    t0 = time.time()
    diarized = _diarized_transcribe(audio_path, num_speakers, language)
    total_s = round(time.time() - t0, 2)
    diarize_s = diarized.get("diarize_s", 0.0)

    meta = {
        "text": diarized["text"],
        "language": diarized["language"],
        "duration": diarized["duration"],
        "speaker_count": diarized["speaker_count"],
        "timing": {
            "transcribe_s": round(total_s - diarize_s, 2),
            "diarize_s": diarize_s,
            "total_s": total_s,
        },
    }
    if diarized.get("warning"):
        meta["warning"] = diarized["warning"]
    return _persist_transcript(diarized["segments"], meta, source_name)


def _run_plain_pipeline(
    audio_path: str,
    language: str | None,
    source_name: str = "audio",
) -> dict:
    """Synchronous plain pipeline (Parakeet, no diarization). Persists the same
    three sidecars as the diarized path. Caller wraps in asyncio.to_thread."""
    t0 = time.time()
    tr = _transcribe(audio_path, language)
    meta = {
        "text": tr["text"],
        "language": tr["language"],
        "duration": tr["duration"],
        "timing": {"transcribe_s": round(time.time() - t0, 2)},
    }
    return _persist_transcript(tr["segments"], meta, source_name)


_AUDIO_EXTENSIONS = {".mp3", ".wav", ".m4a", ".ogg", ".flac", ".webm", ".aac", ".mp4"}
_PLACEHOLDER_RE = re.compile(r"^[<\[{].*[>\]}]$")  # matches <file_id>, [filename], {arg}, etc.


def _latest_audio_upload() -> Path | None:
    """Return the most recently modified audio file in the uploads directory."""
    uploads = get_uploads_dir()
    candidates = [
        f for f in uploads.iterdir() if f.is_file() and f.suffix.lower() in _AUDIO_EXTENSIONS
    ]
    if not candidates:
        return None
    return max(candidates, key=lambda f: f.stat().st_mtime)


def _resolve_audio_input(file: str) -> tuple[Path | None, str]:
    """Resolve a tool input to an absolute path.

    Accepts:
      - OWUI file ID (e.g., 'abc-123' or 'abc-123_meeting.mp3')
      - Filename in the uploads directory
      - Absolute path on the host filesystem
      - Empty string or template placeholder → auto-detects most recent upload

    Returns (path, source_name) or (None, error_message).
    """
    # Empty or literal template placeholder (e.g. '<file_id_or_filename>') →
    # OWUI doesn't surface file references to the pipeline, so auto-detect
    # the most recently uploaded audio file.
    if not file or _PLACEHOLDER_RE.match(file.strip()):
        p = _latest_audio_upload()
        if p is not None:
            logger.info("Auto-detected most recent audio upload: %s", p.name)
            return p, p.name
        return None, "no audio file found in uploads directory — please upload an audio file first"

    # Looks like an absolute path
    if file.startswith("/") or file.startswith("~"):
        p = Path(file).expanduser().resolve()
        if p.is_file():
            return p, p.name
        return None, f"file not found at path: {file}"

    # Try resolving as an upload reference (id or filename)
    p = resolve_upload_path(file)
    if p is not None:
        return p, p.name

    # Last resort: fall back to most recent audio upload with a warning
    fallback = _latest_audio_upload()
    if fallback is not None:
        logger.warning(
            "Could not resolve %r — falling back to most recent upload: %s", file, fallback.name
        )
        return fallback, fallback.name
    return None, f"upload not found: {file!r} (no match in uploads directory)"


# ── FastAPI app ────────────────────────────────────────────────────────────────

# MCP is defined later (needs the tool functions), then the session manager
# lifespan is wired in here so mounted sub-apps get their task group initialized.
from contextlib import asynccontextmanager  # noqa: E402


@asynccontextmanager
async def _lifespan(app: FastAPI):
    async with _mcp_session_manager.run():
        yield


app = FastAPI(title="Portal 5 MLX Transcribe", version="1.0.0", lifespan=_lifespan)


@app.get("/health")
async def health() -> JSONResponse:
    return JSONResponse(
        {
            "status": "ok",
            "service": "mlx-transcribe",
            "asr_model": PARAKEET_MODEL,
            "diarization_model": DIARIZE_MODEL,
            "parakeet_loaded": _parakeet_model is not None,
            "diarizer_loaded": _diarizer_model is not None,
        }
    )


@app.post("/tools/{tool_name}")
async def invoke_tool(tool_name: str, request: Request) -> JSONResponse:
    """REST dispatch endpoint used by portal-pipeline tool_registry."""
    try:
        body = await request.json()
    except Exception:
        body = {}
    arguments = body.get("arguments", {})

    if tool_name == "transcribe_with_speakers":
        file_arg = arguments.get("file", "")
        num_speakers = arguments.get("num_speakers")
        language = arguments.get("language")
        path, source_name = _resolve_audio_input(file_arg)
        if path is None:
            return JSONResponse({"error": source_name})
        async with _pipeline_lock:
            try:
                result = await asyncio.to_thread(
                    _run_pipeline, str(path), language, num_speakers, source_name
                )
                result["engine"] = "parakeet+sortformer"
                return JSONResponse(result)
            except Exception as e:
                logger.error("invoke_tool transcribe_with_speakers failed: %s", e, exc_info=True)
                return JSONResponse({"error": str(e)}, status_code=500)
    elif tool_name == "transcribe_audio":
        file_arg = arguments.get("file", arguments.get("audio_path", ""))
        language = arguments.get("language")
        path, source_name = _resolve_audio_input(file_arg)
        if path is None:
            return JSONResponse({"error": source_name})
        async with _pipeline_lock:
            try:
                result = await asyncio.to_thread(
                    _run_plain_pipeline, str(path), language, source_name
                )
                result["engine"] = "parakeet-tdt-v3"
                return JSONResponse(result)
            except Exception as e:
                logger.error("invoke_tool transcribe_audio failed: %s", e, exc_info=True)
                return JSONResponse({"error": str(e)}, status_code=500)
    else:
        return JSONResponse({"error": f"Unknown tool: {tool_name}"}, status_code=404)


@app.get("/tools")
async def list_tools() -> JSONResponse:
    return JSONResponse(
        {
            "tools": [
                {
                    "name": "transcribe_audio",
                    "description": (
                        "Fast, accurate transcription (Parakeet-TDT-v3, word-level timestamps, "
                        "no speaker labels). Writes JSON + Markdown + Word (.docx) sidecars and "
                        "returns their download URLs. Call with no arguments to auto-detect the "
                        "most recently uploaded audio file."
                    ),
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "file": {
                                "type": "string",
                                "description": "Audio file reference: OWUI file ID, filename in uploads/, or absolute path. Omit to auto-detect most recent upload.",
                            },
                            "language": {
                                "type": "string",
                                "description": "ISO language code (e.g. 'en'). Auto-detected if omitted.",
                            },
                        },
                        "required": [],
                    },
                },
                {
                    "name": "transcribe_with_speakers",
                    "description": (
                        "Transcribe an audio file and label who is speaking: full transcript "
                        "(Parakeet) plus speaker turns (Sortformer diarization), merged at the "
                        "word level. A monologue comes back as one speaker; a conversation gets "
                        "SPEAKER_00/SPEAKER_01/... per turn. Up to 4 speakers, no HuggingFace "
                        "token. Writes JSON + Markdown + Word (.docx) sidecars and returns their "
                        "download URLs. Call with no arguments to auto-detect the most recently "
                        "uploaded audio file."
                    ),
                    "parameters": {
                        "type": "object",
                        "properties": {
                            "file": {
                                "type": "string",
                                "description": "Audio file reference: OWUI file ID, filename in uploads/, or absolute path. Omit to auto-detect most recent upload.",
                            },
                            "num_speakers": {
                                "type": "integer",
                                "description": "Optional cap on speaker count — folds over-segmented speakers into the nearest kept one. Diarization infers the count when omitted.",
                            },
                            "language": {
                                "type": "string",
                                "description": "ISO language code (e.g. 'en'). Auto-detected if omitted.",
                            },
                        },
                        "required": [],
                    },
                },
            ]
        }
    )


@app.get("/files/{filename}", response_model=None)
async def serve_file(filename: str) -> FileResponse | JSONResponse:
    """Serve generated transcript artifacts for browser download."""
    if "/" in filename or ".." in filename:
        return JSONResponse({"error": "invalid filename"}, status_code=400)
    out_dir = get_generated_dir("transcripts")
    path = out_dir / filename
    if not path.exists():
        return JSONResponse({"error": "not found"}, status_code=404)
    return FileResponse(path)


@app.post("/v1/audio/transcribe-with-speakers")
async def http_transcribe(
    file: UploadFile = File(...),  # noqa: B008
    language: str = Form(default="auto"),  # noqa: B008
    num_speakers: int | None = Form(default=None),  # noqa: B008
) -> JSONResponse:
    """Direct HTTP entry point (curl, scripts, batch jobs).

    Accepts a multipart upload, processes synchronously, returns JSON.
    For OWUI integration, use the MCP tool instead.
    """
    tmp_path: str | None = None
    try:
        contents = await file.read()
        suffix = ".wav"
        fname = file.filename or ""
        for ext in [".webm", ".ogg", ".mp4", ".m4a", ".wav", ".mp3", ".flac"]:
            if fname.lower().endswith(ext):
                suffix = ext
                break
        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(contents)
            tmp_path = tmp.name
    except Exception as e:
        return JSONResponse({"error": f"upload failed: {e}"}, status_code=400)

    try:
        async with _pipeline_lock:
            lang = None if language == "auto" else language
            source_name = file.filename or "upload.wav"
            result = await asyncio.to_thread(
                _run_pipeline, tmp_path, lang, num_speakers, source_name
            )
        return JSONResponse(result)
    except Exception as e:
        logger.error("Transcription failed: %s", e, exc_info=True)
        return JSONResponse({"error": str(e)}, status_code=500)
    finally:
        if tmp_path:
            with contextlib.suppress(Exception):
                os.unlink(tmp_path)


# ── MCP wrapper ────────────────────────────────────────────────────────────────

from mcp.server import MCPServer  # noqa: E402

mcp = MCPServer("mlx-transcribe")


@mcp.tool()
async def transcribe_audio(file: str = "", language: str | None = None) -> dict:
    """
    Fast transcription (Parakeet-TDT-v3, word-level timestamps, no speaker labels).

    Saves JSON + Markdown + Word (.docx) sidecars to the workspace
    generated/transcripts/ directory; the full markdown and the download URLs
    (json_url, md_url, docx_url) are in the response.

    Args:
        file: Audio reference. Omit to auto-detect the most recent upload; otherwise an
              OWUI file ID, a filename in uploads/, or an absolute host path.
        language: ISO language code (e.g. 'en'). Auto-detected if omitted.
    """
    path, source_name = _resolve_audio_input(file)
    if path is None:
        return {"error": source_name}
    async with _pipeline_lock:
        try:
            result = await asyncio.to_thread(_run_plain_pipeline, str(path), language, source_name)
            result["engine"] = "parakeet-tdt-v3"
            result["source"] = source_name
            return result
        except Exception as e:
            logger.error("MCP transcribe_audio failed: %s", e, exc_info=True)
            return {"error": str(e)}


@mcp.tool()
async def transcribe_with_speakers(
    file: str = "",
    num_speakers: int | None = None,
    language: str | None = None,
) -> dict:
    """
    Transcribe an audio file and label who is speaking.

    Full transcript via Parakeet-TDT-v3 + speaker turns via Sortformer diarization,
    merged at the word level. A monologue returns one speaker; a conversation
    returns SPEAKER_00/SPEAKER_01/... per turn (up to 4). No HuggingFace token. If
    diarization is skipped (very long file) or fails, the transcript still comes
    back as a single speaker with a ``warning`` — it never truncates.

    Saves JSON + Markdown + Word (.docx) to the workspace generated/transcripts/
    directory; the full markdown and all three download URLs are in the response.

    Args:
        file: Audio file reference. Omit (or leave empty) to auto-detect the
              most recently uploaded audio file. Otherwise accepts:
              - OWUI file ID from a chat attachment (e.g., 'abc-123-def')
              - Filename in the uploads directory (e.g., 'meeting.mp3')
              - Absolute path on the host (e.g., '/Users/me/audio.wav')
        num_speakers: Optional cap on the speaker count (folds over-segmented
              speakers into the nearest kept one). Inferred when omitted.
        language: ISO language code (e.g., 'en'). Auto-detected if omitted.

    Returns:
        dict with:
          - text: full transcript (no speaker labels)
          - language, duration, speaker_count
          - segments: list of {start, end, speaker, text}
          - markdown: full speaker-labeled markdown content (ready to display)
          - json_path, md_path, docx_path: workspace file paths
          - json_url, md_url, docx_url: download URLs (port :8924)
          - timing: {transcribe_s, diarize_s, total_s}
          - engine: "parakeet+sortformer"
          - warning: present only if diarization was skipped or failed

        On error: {"error": "..."}
    """
    logger.info(
        "transcribe_with_speakers called: file=%r num_speakers=%r language=%r",
        file,
        num_speakers,
        language,
    )
    path, source_name = _resolve_audio_input(file)
    if path is None:
        logger.warning("file not resolved: %r — %s", file, source_name)
        return {"error": source_name}

    async with _pipeline_lock:
        try:
            result = await asyncio.to_thread(
                _run_pipeline, str(path), language, num_speakers, source_name
            )
            result["engine"] = "parakeet+sortformer"
            return result
        except Exception as e:
            logger.error("MCP transcribe_with_speakers failed: %s", e, exc_info=True)
            return {"error": str(e)}


# Initialize session manager (must happen after tool definitions, before app start).
# streamable_http_app() lazily creates _session_manager; we surface it so the
# parent lifespan (_lifespan above) can run the task group.
_mcp_sub_app = mcp.streamable_http_app(host=HOST, streamable_http_path="/")
_mcp_session_manager = mcp.session_manager  # noqa: F821 — defined above
app.mount("/mcp", _mcp_sub_app)


# ── Entrypoint ─────────────────────────────────────────────────────────────────


if __name__ == "__main__":
    logger.info("Starting mlx-transcribe on %s:%d", HOST, PORT)
    logger.info("ASR: %s", PARAKEET_MODEL)
    logger.info("Diarizer: %s (skipped above %.0f min)", DIARIZE_MODEL, DIARIZE_MAX_S / 60)
    logger.info("Output dir: %s", get_generated_dir("transcripts"))
    logger.info("Models load lazily on first request. No HF token required.")
    uvicorn.run(app, host=HOST, port=PORT, log_level="info")
